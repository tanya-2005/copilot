"""
exporters.py — turns generated content into real files (.docx) saved to
the generated/ folder, and returns the path so it can be stored on the
GeneratedDocument row.
"""
import io
import json
import os
import re
from docx import Document
from docx.shared import Pt

GENERATED_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "generated")


def _safe_slug(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_").lower()[:60]


def export_tailored_resume_docx(base_resume_text: str, tailored: dict, company: str, title: str) -> str:
    """Produces a simple, clean, ATS-friendly .docx: summary line + tailored
    bullets on top, followed by the full master resume text as reference
    context. (This is a starting draft for you to review/polish — not
    intended as a fully-designed final resume layout.)"""
    os.makedirs(GENERATED_DIR, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Tailored resume draft", level=1)
    doc.add_paragraph(f"For: {title} at {company}")

    if tailored.get("summary_line"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(tailored["summary_line"])

    if tailored.get("bullets"):
        doc.add_heading("Suggested tailored bullets (review before using)", level=2)
        for b in tailored["bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    if tailored.get("ats_keywords"):
        doc.add_heading("ATS keywords this posting looks for", level=2)
        doc.add_paragraph(", ".join(tailored["ats_keywords"]))

    doc.add_heading("Full master resume (reference)", level=2)
    for line in base_resume_text.split("\n"):
        doc.add_paragraph(line)

    filename = f"resume_{_safe_slug(company)}_{_safe_slug(title)}.docx"
    path = os.path.join(GENERATED_DIR, filename)
    doc.save(path)
    return path


def export_cover_letter_docx(letter_text: str, company: str, title: str) -> str:
    os.makedirs(GENERATED_DIR, exist_ok=True)
    doc = Document()
    for para in letter_text.split("\n\n"):
        doc.add_paragraph(para)
    filename = f"cover_letter_{_safe_slug(company)}_{_safe_slug(title)}.docx"
    path = os.path.join(GENERATED_DIR, filename)
    doc.save(path)
    return path


def build_docx_bytes(doc_type: str, content_text: str, company: str, title: str) -> bytes:
    """Rebuilds a .docx in memory from stored content_text — never reads a
    file path from a previous (ephemeral) run. Shared by the dashboard and
    the API, since both only ever have the DB's content_text to work from."""
    doc = Document()
    if doc_type == "resume":
        tailored = json.loads(content_text)
        doc.add_heading("Tailored resume draft", level=1)
        doc.add_paragraph(f"For: {title} at {company}")
        if tailored.get("summary_line"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(tailored["summary_line"])
        if tailored.get("bullets"):
            doc.add_heading("Suggested tailored bullets (review before using)", level=2)
            for b in tailored["bullets"]:
                doc.add_paragraph(b, style="List Bullet")
        if tailored.get("ats_keywords"):
            doc.add_heading("ATS keywords this posting looks for", level=2)
            doc.add_paragraph(", ".join(tailored["ats_keywords"]))
    elif doc_type == "cover_letter":
        for para in content_text.split("\n\n"):
            doc.add_paragraph(para)
    else:
        doc.add_paragraph(content_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
